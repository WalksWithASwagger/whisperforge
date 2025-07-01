# WhisperForge Simple - Clean, Focused Audio Content Platform
import streamlit as st
import os
import tempfile
import time
from datetime import datetime
from typing import Dict, Optional

# Essential imports only
from dotenv import load_dotenv
load_dotenv()

# Page config first
st.set_page_config(
    page_title="WhisperForge",
    page_icon="🌌",
    layout="wide"
)

# Core imports
from core.content_generation import transcribe_audio, generate_wisdom, generate_outline, generate_article, generate_social_content
from core.styling import apply_aurora_theme, create_aurora_header, create_aurora_progress_card, create_aurora_step_card, create_aurora_content_card, AuroraComponents
from core.supabase_integration import get_supabase_client
from core.file_upload import EnhancedLargeFileProcessor

# Apply beautiful theme
apply_aurora_theme()

# === PROMPT LOADING SYSTEM ===
def load_custom_prompts():
    """Load custom prompts from the prompts directory"""
    prompts = {}
    prompt_dir = "prompts/default"
    
    if os.path.exists(prompt_dir):
        for filename in os.listdir(prompt_dir):
            if filename.endswith('.md'):
                prompt_name = filename.replace('.md', '')
                try:
                    with open(os.path.join(prompt_dir, filename), 'r', encoding='utf-8') as f:
                        prompts[prompt_name] = f.read()
                except Exception as e:
                    st.warning(f"Failed to load prompt {filename}: {e}")
    
    return prompts

def load_template(template_name: str) -> Optional[str]:
    """Load an article template by name"""
    template_path = os.path.join('templates', f'{template_name}.md')
    if os.path.exists(template_path):
        return open(template_path, 'r', encoding='utf-8').read()
    return None

def get_prompt_for_step(step_name: str, custom_prompts: Dict[str, str] = None) -> Optional[str]:
    """Get the appropriate prompt for a pipeline step"""
    if not custom_prompts:
        custom_prompts = load_custom_prompts()
    
    # Map step names to prompt files
    prompt_mapping = {
        'wisdom': 'wisdom_extraction',
        'outline': 'outline_creation', 
        'social': 'social_media',
        'article': 'article_generation'  # We'll create this
    }
    
    prompt_key = prompt_mapping.get(step_name)
    if prompt_key and prompt_key in custom_prompts:
        return custom_prompts[prompt_key]
    
    return None

# === NOTION INTEGRATION ===
def create_notion_page(title: str, content_data: Dict[str, str]) -> Optional[str]:
    """Create a Notion page with WhisperForge content"""
    try:
        from notion_client import Client
        
        api_key = os.getenv("NOTION_API_KEY")
        database_id = os.getenv("NOTION_DATABASE_ID")
        
        if not api_key or not database_id:
            st.warning("⚠️ Notion not configured. Set NOTION_API_KEY and NOTION_DATABASE_ID to auto-publish.")
            return None
        
        client = Client(auth=api_key)
        
        # Build content blocks
        children = []
        
        # Add beautiful header with summary
        children.append({
            "type": "heading_1",
            "heading_1": {
                "rich_text": [
                    {"type": "text", "text": {"content": "🌌 "}, "annotations": {"color": "blue"}},
                    {"type": "text", "text": {"content": title}, "annotations": {"bold": True}}
                ]
            }
        })
        
        # Add creation info
        children.append({
            "type": "paragraph",
            "paragraph": {
                "rich_text": [
                    {"type": "text", "text": {"content": "✨ Generated with "}},
                    {"type": "text", "text": {"content": "WhisperForge Aurora"}, 
                     "annotations": {"bold": True, "color": "blue"}},
                    {"type": "text", "text": {"content": f" • {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"}}
                ]
            }
        })
        
        children.append({"type": "divider", "divider": {}})
        
        # Add wisdom summary callout if exists
        if content_data.get('wisdom'):
            children.append({
                "type": "callout",
                "callout": {
                    "rich_text": [
                        {"type": "text", "text": {"content": "Key Insights & Wisdom"}},
                        {"type": "text", "text": {"content": f"\n\n{content_data['wisdom'][:1800]}"}}
                    ],
                    "color": "purple_background",
                    "icon": {"type": "emoji", "emoji": "💡"}
                }
            })
        
        # Add content sections as toggles
        sections = [
            ("📝 Transcript", content_data.get('transcript')),
            ("💡 Wisdom", content_data.get('wisdom')),
            ("🔍 Research Links", content_data.get('research')),
            ("📋 Outline", content_data.get('outline')),
            ("📰 Article", content_data.get('article')),
            ("📱 Social Content", content_data.get('social_content'))
        ]
        
        for section_title, section_content in sections:
            if section_content:
                # Handle research data specially
                if section_title == "🔍 Research Links" and isinstance(section_content, dict):
                    research_children = []
                    entities = section_content.get('entities', [])
                    
                    if entities:
                        for entity in entities[:5]:  # Limit entities
                            entity_name = entity.get('name', 'Unknown Entity')
                            why_matters = entity.get('why_matters', 'No description available')
                            links = entity.get('links', [])
                            
                            # Entity as beautiful callout
                            research_children.append({
                                "type": "callout",
                                "callout": {
                                    "rich_text": [
                                        {"type": "text", "text": {"content": entity_name}, "annotations": {"bold": True}},
                                        {"type": "text", "text": {"content": f"\n{why_matters}"}}
                                    ],
                                    "color": "blue_background",
                                    "icon": {"type": "emoji", "emoji": "🔬"}
                                }
                            })
                            
                            # Links as bulleted list
                            if links:
                                for link in links[:3]:  # Limit links
                                    link_title = link.get('title', 'Link')
                                    link_url = link.get('url', '#')
                                    link_desc = link.get('description', '')
                                    is_gem = link.get('is_gem', False)
                                    
                                    gem_icon = "💎" if is_gem else "🔗"
                                    color = "orange" if is_gem else "default"
                                    
                                    research_children.append({
                                        "type": "bulleted_list_item",
                                        "bulleted_list_item": {
                                            "rich_text": [
                                                {"type": "text", "text": {"content": f"{gem_icon} "}, "annotations": {"color": color}},
                                                {"type": "text", "text": {"content": link_title}, "annotations": {"bold": True}},
                                                {"type": "text", "text": {"content": f" - {link_desc}"}, "annotations": {"italic": True}}
                                            ]
                                        }
                                    })
                    else:
                        research_children.append({
                            "type": "paragraph",
                            "paragraph": {
                                "rich_text": [{"type": "text", "text": {"content": "No research entities found."}}]
                            }
                        })
                    
                    children.append({
                        "type": "toggle",
                        "toggle": {
                            "rich_text": [{"type": "text", "text": {"content": section_title}}],
                            "children": research_children
                        }
                    })
                else:
                    # Handle regular text content
                    if isinstance(section_content, str):
                        # Chunk content for Notion's limits
                        chunks = [section_content[i:i+1800] for i in range(0, len(section_content), 1800)]
                        
                        children.append({
                            "type": "toggle",
                            "toggle": {
                                "rich_text": [{"type": "text", "text": {"content": section_title}}],
                                "children": [
                                    {
                                        "type": "paragraph",
                                        "paragraph": {
                                            "rich_text": [{"type": "text", "text": {"content": chunk}}]
                                        }
                                    } for chunk in chunks[:5]  # Limit chunks
                                ]
                            }
                        })
        
        # Add beautiful footer
        children.extend([
            {"type": "divider", "divider": {}},
            {
                "type": "callout",
                "callout": {
                    "rich_text": [
                        {"type": "text", "text": {"content": "Content Generation Complete"}, "annotations": {"bold": True}},
                        {"type": "text", "text": {"content": f"\n\n🤖 AI Pipeline: 8 steps completed successfully"}},
                        {"type": "text", "text": {"content": f"\n⏱️ Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"}},
                        {"type": "text", "text": {"content": f"\n🌌 Powered by WhisperForge Aurora"}}
                    ],
                    "color": "green_background",
                    "icon": {"type": "emoji", "emoji": "✅"}
                }
            }
        ])
        
        # Create the page
        response = client.pages.create(
            parent={"database_id": database_id},
            icon={"type": "emoji", "emoji": "🌌"},
            properties={
                "Name": {"title": [{"text": {"content": title[:100]}}]}
            },
            children=children[:50]  # Limit total blocks
        )
        
        if response and 'id' in response:
            page_id = response['id']
            page_url = f"https://notion.so/{page_id.replace('-', '')}"
            return page_url
        
        return None
        
    except ImportError:
        st.warning("⚠️ Install notion-client to enable Notion publishing: pip install notion-client")
        return None
    except Exception as e:
        st.error(f"❌ Notion publishing failed: {str(e)}")
        return None

def generate_ai_title(transcript: str) -> str:
    """Generate an AI title for the content"""
    try:
        from core.content_generation import generate_content
        
        prompt = f"""Generate a concise, descriptive title (max 60 characters) for this audio transcript:

{transcript[:500]}...

Title should be:
- Clear and specific
- Professional
- Capture the main topic
- No quotes or special characters

Title:"""
        
        title = generate_content(prompt, "OpenAI", "gpt-4", {})
        return title.strip().replace('"', '').replace("'", "")[:60]
    except:
        return f"WhisperForge Content - {datetime.now().strftime('%Y-%m-%d %H:%M')}"

# === SIMPLE AUTHENTICATION ===
def init_session():
    """Initialize simple session state"""
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'user_id' not in st.session_state:
        st.session_state.user_id = None
    if 'user_email' not in st.session_state:
        st.session_state.user_email = None

def show_login():
    """Simple test login"""
    create_aurora_header()
    
    st.markdown("### 🔐 Login to WhisperForge")
    
    # Test login button
    if st.button("🚀 Login with Test Account", type="primary", use_container_width=True):
        st.session_state.authenticated = True
        st.session_state.user_id = 1
        st.session_state.user_email = "test@whisperforge.ai"
        st.success("✅ Logged in successfully!")
        time.sleep(1)
        st.rerun()
    
    st.markdown("---")
    st.markdown("**Demo Mode**: Click above to access WhisperForge")

# === CORE PROCESSING PIPELINE ===

def show_processing_pipeline(current_step=0, step_progress=0, total_progress=0, status_message="", processing_time=""):
    """Display beautiful Aurora-styled processing pipeline visualization"""
    
    # Define the 6-step pipeline
    pipeline_steps = [
        {
            "icon": "🎤",
            "title": "Transcription",
            "description": "Converting audio to text using Whisper AI",
            "status": "pending"
        },
        {
            "icon": "💡",
            "title": "Wisdom Extraction",
            "description": "Extracting key insights and wisdom",
            "status": "pending"
        },
        {
            "icon": "📋",
            "title": "Outline Creation",
            "description": "Structuring content with clear outline",
            "status": "pending"
        },
        {
            "icon": "📰",
            "title": "Article Generation",
            "description": "Creating comprehensive article content",
            "status": "pending"
        },
        {
            "icon": "📱",
            "title": "Social Content",
            "description": "Generating social media posts",
            "status": "pending"
        },
        {
            "icon": "🌌",
            "title": "Notion Publishing",
            "description": "Publishing to Notion workspace",
            "status": "pending"
        }
    ]
    
    # Update step statuses based on current progress
    for i, step in enumerate(pipeline_steps):
        if i < current_step:
            step["status"] = "completed"
        elif i == current_step:
            step["status"] = "active"
        else:
            step["status"] = "pending"
    
    # Create the pipeline visualization HTML
    steps_html = ""
    for i, step in enumerate(pipeline_steps):
        progress_width = step_progress if i == current_step else (100 if step["status"] == "completed" else 0)
        
        status_text = {
            "pending": "Waiting",
            "active": "Processing",
            "completed": "Complete",
            "error": "Error"
        }.get(step["status"], "Waiting")
        
        steps_html += f"""
        <div class="aurora-pipeline-step {step['status']}">
            <div class="aurora-step-progress" style="width: {progress_width}%;"></div>
            <span class="aurora-step-icon">{step['icon']}</span>
            <h4 class="aurora-step-title">{step['title']}</h4>
            <p class="aurora-step-description">{step['description']}</p>
            <div class="aurora-step-status">{status_text}</div>
        </div>
        """
    
    # Create the complete pipeline HTML
    pipeline_html = f"""
    <div class="aurora-pipeline-container">
        <div class="aurora-pipeline-header">
            <h3 class="aurora-pipeline-title">Content Transformation Pipeline</h3>
            <p class="aurora-pipeline-subtitle">6-Step AI-Powered Processing</p>
        </div>
        
        <div class="aurora-overall-progress">
            <div class="aurora-progress-header">
                <span class="aurora-progress-label">Overall Progress</span>
                <span class="aurora-progress-percentage">{total_progress}%</span>
            </div>
            <div class="aurora-progress-bar-container">
                <div class="aurora-progress-bar" style="width: {total_progress}%;"></div>
            </div>
        </div>
        
        <div class="aurora-pipeline-steps">
            {steps_html}
        </div>
        
        {f'''
        <div class="aurora-processing-status active">
            <span class="aurora-status-icon">⚡</span>
            <span class="aurora-status-text">{status_message}</span>
            <span class="aurora-status-time">{processing_time}</span>
        </div>
        ''' if status_message else ''}
    </div>
    """
    
    st.markdown(pipeline_html, unsafe_allow_html=True)

def process_audio_pipeline(audio_file):
    """Core audio to content pipeline with beautiful Aurora visualization"""
    import time
    from datetime import datetime
    
    results = {}
    start_time = time.time()
    
    # Load custom prompts
    custom_prompts = load_custom_prompts()
    if custom_prompts:
        st.info(f"📝 Using {len(custom_prompts)} custom prompts")
    
    # Initialize beautiful pipeline visualization
    pipeline_placeholder = st.empty()
    
    # Create real-time content display containers
    st.markdown("### 🌌 Live Content Generation")
    
    # Create expandable containers for each step
    transcript_container = st.expander("🎙️ Transcription", expanded=False)
    wisdom_container = st.expander("💡 Wisdom Extraction", expanded=False)
    outline_container = st.expander("📋 Outline Creation", expanded=False)
    article_container = st.expander("📝 Article Generation", expanded=False)
    social_container = st.expander("📱 Social Content", expanded=False)
    notion_container = st.expander("🌌 Notion Publishing", expanded=False)
    
    try:
        # Step 1: Transcription
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=0, 
                step_progress=0, 
                total_progress=0,
                status_message="Starting transcription process...",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        # Import transcription function
        from core.content_generation import transcribe_audio
        
        # Create temporary file
        import tempfile
        import os
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(audio_file.name)[1]) as tmp_file:
            tmp_file.write(audio_file.getvalue())
            tmp_file_path = tmp_file.name
        
        try:
            # Transcription with progress updates
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=0, 
                    step_progress=50, 
                    total_progress=8,
                    status_message="Transcribing audio with Whisper AI...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            transcript = transcribe_audio(tmp_file_path)
            if not transcript or "Error" in transcript:
                st.error(f"Transcription failed: {transcript}")
                return None
            
            results['transcript'] = transcript
            
            # Stream transcript to UI immediately
            with transcript_container:
                st.markdown("**✅ Transcription Complete**")
                st.text_area("Transcript", transcript, height=200, disabled=True)
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=0, 
                    step_progress=100, 
                    total_progress=17,
                    status_message="Transcription complete!",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            # Step 2: Wisdom Extraction
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=1, 
                    step_progress=0, 
                    total_progress=17,
                    status_message="Extracting wisdom and insights...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            from core.content_generation import generate_wisdom
            wisdom_prompt = get_prompt_for_step('wisdom', custom_prompts)
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=1, 
                    step_progress=50, 
                    total_progress=25,
                    status_message="Analyzing content for key insights...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            wisdom = generate_wisdom(transcript, custom_prompt=wisdom_prompt, knowledge_base={})
            results['wisdom'] = wisdom
            
            # Stream wisdom to UI immediately
            with wisdom_container:
                st.markdown("**✅ Wisdom Extraction Complete**")
                st.markdown(wisdom)
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=1, 
                    step_progress=100, 
                    total_progress=33,
                    status_message="Wisdom extraction complete!",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            # Step 3: Outline Creation
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=2, 
                    step_progress=0, 
                    total_progress=33,
                    status_message="Creating structured outline...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            from core.content_generation import generate_outline
            outline_prompt = get_prompt_for_step('outline', custom_prompts)
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=2, 
                    step_progress=50, 
                    total_progress=42,
                    status_message="Structuring content hierarchy...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            outline = generate_outline(transcript, wisdom, custom_prompt=outline_prompt, knowledge_base={})
            results['outline'] = outline
            
            # Stream outline to UI immediately
            with outline_container:
                st.markdown("**✅ Outline Creation Complete**")
                st.markdown(outline)
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=2, 
                    step_progress=100, 
                    total_progress=50,
                    status_message="Outline creation complete!",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            # Step 4: Article Generation
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=3, 
                    step_progress=0, 
                    total_progress=50,
                    status_message="Generating comprehensive article...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            from core.content_generation import generate_article
            article_prompt = get_prompt_for_step('article', custom_prompts)
            selected_template = st.session_state.get('article_template')
            if selected_template:
                template_text = load_template(selected_template)
                if template_text:
                    article_prompt = template_text + "\n" + article_prompt
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=3, 
                    step_progress=50, 
                    total_progress=58,
                    status_message="Writing detailed article content...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            article = generate_article(transcript, wisdom, outline, custom_prompt=article_prompt, knowledge_base={})
            results['article'] = article
            
            # Stream article to UI immediately
            with article_container:
                st.markdown("**✅ Article Generation Complete**")
                st.markdown(article)
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=3, 
                    step_progress=100, 
                    total_progress=67,
                    status_message="Article generation complete!",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            # Step 5: Social Content
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=4, 
                    step_progress=0, 
                    total_progress=67,
                    status_message="Creating social media content...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            from core.content_generation import generate_social_content
            social_prompt = get_prompt_for_step('social', custom_prompts)
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=4, 
                    step_progress=50, 
                    total_progress=75,
                    status_message="Generating social media posts...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            social = generate_social_content(wisdom, outline, article, custom_prompt=social_prompt, knowledge_base={})
            results['social_content'] = social
            
            # Stream social content to UI immediately
            with social_container:
                st.markdown("**✅ Social Content Creation Complete**")
                st.markdown(social)
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=4, 
                    step_progress=100, 
                    total_progress=83,
                    status_message="Social content creation complete!",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            # Step 6: Auto-publish to Notion
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=5, 
                    step_progress=0, 
                    total_progress=83,
                    status_message="Publishing to Notion workspace...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            if os.getenv("NOTION_API_KEY") and os.getenv("NOTION_DATABASE_ID"):
                # Generate AI title
                ai_title = generate_ai_title(transcript)
                
                with pipeline_placeholder.container():
                    show_processing_pipeline(
                        current_step=5, 
                        step_progress=30, 
                        total_progress=88,
                        status_message="Creating Notion page structure...",
                        processing_time=f"{time.time() - start_time:.1f}s"
                    )
                
                with pipeline_placeholder.container():
                    show_processing_pipeline(
                        current_step=5, 
                        step_progress=60, 
                        total_progress=92,
                        status_message="Uploading content to Notion...",
                        processing_time=f"{time.time() - start_time:.1f}s"
                    )
                
                # Publish to Notion
                notion_url = create_notion_page(ai_title, results)
                if notion_url:
                    results['notion_url'] = notion_url
                    
                    # Stream Notion success to UI
                    with notion_container:
                        st.markdown("**✅ Notion Publishing Complete**")
                        st.markdown(f"**Page Title:** {ai_title}")
                        st.markdown(f"🔗 [Open in Notion]({notion_url})")
                else:
                    # Stream Notion failure to UI
                    with notion_container:
                        st.markdown("**⚠️ Notion Publishing Failed**")
                        st.warning("Check your Notion API configuration in Settings.")
            else:
                # Show disabled status in UI
                with notion_container:
                    st.markdown("**ℹ️ Notion Publishing Disabled**")
                    st.info("Configure Notion API in Settings to enable auto-publishing.")
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=5, 
                    step_progress=90, 
                    total_progress=96,
                    status_message="Saving to database...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            # Save to Supabase database
            try:
                save_content_to_db(results)
            except Exception as e:
                st.warning(f"⚠️ Content saved locally but database save failed: {e}")
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=5, 
                    step_progress=100, 
                    total_progress=100,
                    status_message="Pipeline complete! All content generated successfully.",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            # Aurora completion celebration
            st.markdown("""
            <div class="aurora-celebration">
                <h1 class="aurora-celebration-title">Pipeline Complete!</h1>
                <p class="aurora-celebration-subtitle">Your content has been transformed with AI magic</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Clear the pipeline display after a moment
            time.sleep(2)
            pipeline_placeholder.empty()
            
            return results
            
        finally:
            # Cleanup temporary file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)
                
    except Exception as e:
        # Show error state
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=0, 
                step_progress=0, 
                total_progress=0,
                status_message=f"Error: {str(e)}",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        st.error(f"Pipeline failed: {str(e)}")
        return None

def process_audio_pipeline_live(audio_file):
    """Run pipeline with StreamingPipelineController"""
    from core.streaming_pipeline import get_pipeline_controller

    controller = get_pipeline_controller()
    controller.start_pipeline(audio_file)
    while controller.process_next_step():
        pass
    return controller.get_results()

def process_audio_pipeline_with_transcript(transcript: str):
    """Process audio pipeline with pre-transcribed content using beautiful Aurora visualization"""
    import time
    from datetime import datetime
    
    results = {'transcript': transcript}
    start_time = time.time()
    
    # Load custom prompts
    custom_prompts = load_custom_prompts()
    if custom_prompts:
        st.info(f"📝 Using {len(custom_prompts)} custom prompts")
    
    # Initialize beautiful pipeline visualization (starting from step 1)
    pipeline_placeholder = st.empty()
    
    # Create real-time content display containers
    st.markdown("### 🌌 Live Content Generation")
    
    # Create expandable containers for each step (skip transcription)
    wisdom_container = st.expander("💡 Wisdom Extraction", expanded=False)
    outline_container = st.expander("📋 Outline Creation", expanded=False)
    article_container = st.expander("📝 Article Generation", expanded=False)
    social_container = st.expander("📱 Social Content", expanded=False)
    notion_container = st.expander("🌌 Notion Publishing", expanded=False)
    
    try:
        # Show initial state with transcription already complete
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=1, 
                step_progress=0, 
                total_progress=17,
                status_message=f"Using pre-transcribed content ({len(transcript)} characters)",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        # Step 2: Wisdom Extraction
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=1, 
                step_progress=0, 
                total_progress=17,
                status_message="Extracting wisdom and insights...",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        from core.content_generation import generate_wisdom
        wisdom_prompt = get_prompt_for_step('wisdom', custom_prompts)
        
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=1, 
                step_progress=50, 
                total_progress=25,
                status_message="Analyzing content for key insights...",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        wisdom = generate_wisdom(transcript, custom_prompt=wisdom_prompt, knowledge_base={})
        results['wisdom'] = wisdom
        
        # Stream wisdom to UI immediately
        with wisdom_container:
            st.markdown("**✅ Wisdom Extraction Complete**")
            st.markdown(wisdom)
        
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=1, 
                step_progress=100, 
                total_progress=33,
                status_message="Wisdom extraction complete!",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        # Step 3: Outline Creation
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=2, 
                step_progress=0, 
                total_progress=33,
                status_message="Creating structured outline...",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        from core.content_generation import generate_outline
        outline_prompt = get_prompt_for_step('outline', custom_prompts)
        
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=2, 
                step_progress=50, 
                total_progress=42,
                status_message="Structuring content hierarchy...",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        outline = generate_outline(transcript, wisdom, custom_prompt=outline_prompt, knowledge_base={})
        results['outline'] = outline
        
        # Stream outline to UI immediately
        with outline_container:
            st.markdown("**✅ Outline Creation Complete**")
            st.markdown(outline)
        
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=2, 
                step_progress=100, 
                total_progress=50,
                status_message="Outline creation complete!",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        # Step 4: Article Generation
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=3, 
                step_progress=0, 
                total_progress=50,
                status_message="Generating comprehensive article...",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        from core.content_generation import generate_article
        article_prompt = get_prompt_for_step('article', custom_prompts)
        selected_template = st.session_state.get('article_template')
        if selected_template:
            template_text = load_template(selected_template)
            if template_text:
                article_prompt = template_text + "\n" + article_prompt
        
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=3, 
                step_progress=50, 
                total_progress=58,
                status_message="Writing detailed article content...",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        article = generate_article(transcript, wisdom, outline, custom_prompt=article_prompt, knowledge_base={})
        results['article'] = article
        
        # Stream article to UI immediately
        with article_container:
            st.markdown("**✅ Article Generation Complete**")
            st.markdown(article)
        
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=3, 
                step_progress=100, 
                total_progress=67,
                status_message="Article generation complete!",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        # Step 5: Social Content
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=4, 
                step_progress=0, 
                total_progress=67,
                status_message="Creating social media content...",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        from core.content_generation import generate_social_content
        social_prompt = get_prompt_for_step('social', custom_prompts)
        
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=4, 
                step_progress=50, 
                total_progress=75,
                status_message="Generating social media posts...",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        social = generate_social_content(wisdom, outline, article, custom_prompt=social_prompt, knowledge_base={})
        results['social_content'] = social
        
        # Stream social content to UI immediately
        with social_container:
            st.markdown("**✅ Social Content Creation Complete**")
            st.markdown(social)
        
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=4, 
                step_progress=100, 
                total_progress=83,
                status_message="Social content creation complete!",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        # Step 6: Auto-publish to Notion
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=5, 
                step_progress=0, 
                total_progress=83,
                status_message="Publishing to Notion workspace...",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        if os.getenv("NOTION_API_KEY") and os.getenv("NOTION_DATABASE_ID"):
            # Generate AI title
            ai_title = generate_ai_title(transcript)
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=5, 
                    step_progress=30, 
                    total_progress=88,
                    status_message="Creating Notion page structure...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            with pipeline_placeholder.container():
                show_processing_pipeline(
                    current_step=5, 
                    step_progress=60, 
                    total_progress=92,
                    status_message="Uploading content to Notion...",
                    processing_time=f"{time.time() - start_time:.1f}s"
                )
            
            # Publish to Notion
            notion_url = create_notion_page(ai_title, results)
            if notion_url:
                results['notion_url'] = notion_url
                
                # Stream Notion success to UI
                with notion_container:
                    st.markdown("**✅ Notion Publishing Complete**")
                    st.markdown(f"**Page Title:** {ai_title}")
                    st.markdown(f"🔗 [Open in Notion]({notion_url})")
            else:
                # Stream Notion failure to UI
                with notion_container:
                    st.markdown("**⚠️ Notion Publishing Failed**")
                    st.warning("Check your Notion API configuration in Settings.")
        else:
            # Show disabled status in UI
            with notion_container:
                st.markdown("**ℹ️ Notion Publishing Disabled**")
                st.info("Configure Notion API in Settings to enable auto-publishing.")
        
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=5, 
                step_progress=90, 
                total_progress=96,
                status_message="Saving to database...",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        # Save to Supabase database
        try:
            save_content_to_db(results)
        except Exception as e:
            st.warning(f"⚠️ Content saved locally but database save failed: {e}")
        
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=5, 
                step_progress=100, 
                total_progress=100,
                status_message="Pipeline complete! All content generated successfully.",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        
        # Aurora completion celebration
        st.markdown("""
        <div class="aurora-celebration">
            <h1 class="aurora-celebration-title">Pipeline Complete!</h1>
            <p class="aurora-celebration-subtitle">Your content has been transformed with AI magic</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Clear the pipeline display after a moment
        time.sleep(2)
        pipeline_placeholder.empty()
        
        return results
        
    except Exception as e:
        # Show error state
        with pipeline_placeholder.container():
            show_processing_pipeline(
                current_step=1, 
                step_progress=0, 
                total_progress=17,
                status_message=f"Error: {str(e)}",
                processing_time=f"{time.time() - start_time:.1f}s"
            )
        st.error(f"Pipeline failed: {str(e)}")
        return None

def save_content_to_db(content_data):
    """Save generated content to database"""
    try:
        db = get_supabase_client()
        if db and st.session_state.user_id:
            content_id = db.save_content(st.session_state.user_id, {
                **content_data,
                'title': f"WhisperForge Content {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                'created_at': datetime.now().isoformat()
            })
            if content_id:
                st.success(f"💾 Content saved to database (ID: {content_id})")
    except Exception as e:
        st.warning(f"Database save failed: {e}")

# === CONTENT DISPLAY ===
def create_enhanced_aurora_content_card(title, content, content_type="text", icon="📄"):
    """Create a beautiful enhanced Aurora content card with copy functionality and animations"""
    import uuid
    
    # Generate unique IDs for this card
    card_id = f"card_{uuid.uuid4().hex[:8]}"
    copy_btn_id = f"copy_{uuid.uuid4().hex[:8]}"
    expand_btn_id = f"expand_{uuid.uuid4().hex[:8]}"
    full_content_id = f"full_{uuid.uuid4().hex[:8]}"
    
    # Calculate content stats
    word_count = len(content.split()) if content else 0
    char_count = len(content) if content else 0
    
    # Determine if content needs truncation
    preview_length = 300
    needs_expansion = len(content) > preview_length
    preview_content = content[:preview_length] + "..." if needs_expansion else content
    
    # Content type specific styling
    type_class = content_type.lower()
    
    # Create the enhanced card HTML
    card_html = f"""
    <div class="aurora-content-card {type_class}" id="{card_id}">
        <div class="aurora-content-card-header">
            <h3 class="aurora-content-card-title">
                <span class="aurora-content-card-icon">{icon}</span>
                {title}
            </h3>
            <div class="aurora-content-card-actions">
                <button class="aurora-content-action-btn copy-btn" id="{copy_btn_id}" onclick="copyToClipboard('{card_id}', '{copy_btn_id}')">
                    <span>📋</span>
                    <span class="copy-text">Copy</span>
                </button>
                <button class="aurora-content-action-btn" onclick="downloadContent('{title}', `{content.replace('`', '\\`')}`)">
                    <span>💾</span>
                    <span>Save</span>
                </button>
            </div>
        </div>
        
        <div class="aurora-content-card-body">
            <div class="aurora-content-preview" id="preview_{card_id}">
                {preview_content}
            </div>
            
            {f'''
            <div class="aurora-content-full" id="{full_content_id}">
                {content[preview_length:]}
            </div>
            ''' if needs_expansion else ''}
            
            <div class="aurora-content-stats">
                <div class="aurora-content-word-count">
                    <span>📊</span>
                    <span>{word_count} words • {char_count} characters</span>
                </div>
                
                {f'''
                <button class="aurora-content-expand-btn" id="{expand_btn_id}" onclick="toggleExpand('{full_content_id}', '{expand_btn_id}')">
                    <span>Show more</span>
                    <span class="aurora-content-expand-icon">▼</span>
                </button>
                ''' if needs_expansion else ''}
            </div>
        </div>
        
        <!-- Hidden content for copying -->
        <textarea id="content_{card_id}" style="position: absolute; left: -9999px; opacity: 0;">{content}</textarea>
    </div>
    
    <script>
    function copyToClipboard(cardId, btnId) {{
        const content = document.getElementById('content_' + cardId).value;
        const button = document.getElementById(btnId);
        const copyText = button.querySelector('.copy-text');
        
        navigator.clipboard.writeText(content).then(function() {{
            // Success feedback
            button.classList.add('copied');
            copyText.textContent = 'Copied!';
            
            // Reset after 2 seconds
            setTimeout(() => {{
                button.classList.remove('copied');
                copyText.textContent = 'Copy';
            }}, 2000);
        }}).catch(function(err) {{
            console.error('Failed to copy: ', err);
            // Fallback for older browsers
            const textarea = document.getElementById('content_' + cardId);
            textarea.select();
            document.execCommand('copy');
            
            button.classList.add('copied');
            copyText.textContent = 'Copied!';
            setTimeout(() => {{
                button.classList.remove('copied');
                copyText.textContent = 'Copy';
            }}, 2000);
        }});
    }}
    
    function toggleExpand(fullContentId, btnId) {{
        const fullContent = document.getElementById(fullContentId);
        const button = document.getElementById(btnId);
        const buttonText = button.querySelector('span:first-child');
        const icon = button.querySelector('.aurora-content-expand-icon');
        
        if (fullContent.classList.contains('expanded')) {{
            fullContent.classList.remove('expanded');
            button.classList.remove('expanded');
            buttonText.textContent = 'Show more';
        }} else {{
            fullContent.classList.add('expanded');
            button.classList.add('expanded');
            buttonText.textContent = 'Show less';
        }}
    }}
    
    function downloadContent(title, content) {{
        const blob = new Blob([content], {{ type: 'text/plain' }});
        const url = window.URL.createObjectURL(blob);
        const a = document.createElement('a');
        a.href = url;
        a.download = title.replace(/[^a-z0-9]/gi, '_').toLowerCase() + '.txt';
        document.body.appendChild(a);
        a.click();
        document.body.removeChild(a);
        window.URL.revokeObjectURL(url);
    }}
    </script>
    """
    
    st.markdown(card_html, unsafe_allow_html=True)

def show_results(results):
    """Display generated content with beautiful Aurora styling and enhanced UX"""
    if not results:
        return
    
    # Aurora header for results with enhanced styling
    st.markdown("""
    <div class="aurora-results-header">
        <h1 class="aurora-results-title">✨ Content Generated Successfully!</h1>
        <p class="aurora-results-subtitle">Your audio has been transformed with AI magic</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Aurora Notion link if available with enhanced styling
    if results.get('notion_url'):
        st.markdown(f"""
        <div class="aurora-notion-link">
            <a href="{results['notion_url']}" target="_blank" class="aurora-notion-button">
                🌌 View in Notion Workspace
            </a>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("---")
    
    # Enhanced content overview stats
    total_words = sum(len(str(results.get(key, '')).split()) for key in ['transcript', 'wisdom', 'outline', 'article', 'social_content'])
    content_types = len([k for k in ['transcript', 'wisdom', 'outline', 'article', 'social_content'] if results.get(k)])
    
    st.markdown(f"""
    <div class="aurora-content-overview">
        <div class="aurora-overview-stats">
            <div class="aurora-stat-item">
                <span class="aurora-stat-icon">📊</span>
                <span class="aurora-stat-value">{total_words:,}</span>
                <span class="aurora-stat-label">Total Words</span>
            </div>
            <div class="aurora-stat-item">
                <span class="aurora-stat-icon">📄</span>
                <span class="aurora-stat-value">{content_types}</span>
                <span class="aurora-stat-label">Content Types</span>
            </div>
            <div class="aurora-stat-item">
                <span class="aurora-stat-icon">⚡</span>
                <span class="aurora-stat-value">AI</span>
                <span class="aurora-stat-label">Generated</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Prepare tab data for custom Aurora tabs
    tab_data = []
    
    # Add main content tabs
    if results.get('transcript'):
        tab_data.append({
            'title': 'Transcript',
            'icon': '🎙️',
            'type': 'transcript',
            'content': results['transcript']
        })
    
    if results.get('wisdom'):
        tab_data.append({
            'title': 'Wisdom',
            'icon': '💎',
            'type': 'wisdom',
            'content': results['wisdom']
        })
    
    if results.get('outline'):
        tab_data.append({
            'title': 'Outline',
            'icon': '📋',
            'type': 'outline',
            'content': results['outline']
        })
    
    if results.get('article'):
        tab_data.append({
            'title': 'Article',
            'icon': '📰',
            'type': 'article',
            'content': results['article']
        })
    
    if results.get('social_content'):
        tab_data.append({
            'title': 'Social',
            'icon': '📱',
            'type': 'social',
            'content': results['social_content']
        })
    
    # Add Editor tab if editor content exists
    if results.get('editor_notes') or results.get('revised_content'):
        editor_content = ""
        
        if results.get('editor_notes'):
            editor_content += "=== EDITOR NOTES ===\n\n"
            for section, notes in results['editor_notes'].items():
                if notes:
                    editor_content += f"## {section.title()} Notes:\n{notes}\n\n"
        
        if results.get('revised_content'):
            editor_content += "\n=== REVISED CONTENT ===\n\n"
            for section, content in results['revised_content'].items():
                if content:
                    editor_content += f"## Revised {section.title()}:\n{content}\n\n"
        
        tab_data.append({
            'title': 'Editor Review',
            'icon': '📝',
            'type': 'editor',
            'content': editor_content
        })
    
    # Display the custom Aurora tabs
    if tab_data:
        create_aurora_tabs(tab_data, default_tab=0)
    else:
        st.warning("No content available to display.")
    
    # Add export all functionality
    st.markdown("---")
    st.markdown("""
    <div class="aurora-export-section">
        <h3 class="aurora-export-title">📦 Additional Export Options</h3>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📄 Export as Text", use_container_width=True):
            export_content = create_text_export(results)
            st.download_button(
                label="💾 Download Text File",
                data=export_content,
                file_name=f"whisperforge_content_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                mime="text/plain",
                use_container_width=True
            )
        if st.button("📝 Export as Markdown", use_container_width=True):
            md_content = export_to_markdown(results)
            st.download_button(
                label="💾 Download Markdown",
                data=md_content,
                file_name=f"whisperforge_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                mime="text/markdown",
                use_container_width=True
            )

    with col2:
        if st.button("📊 Export as JSON", use_container_width=True):
            import json
            json_content = json.dumps(results, indent=2, ensure_ascii=False)
            st.download_button(
                label="💾 Download JSON File",
                data=json_content,
                file_name=f"whisperforge_data_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json",
                use_container_width=True
            )
        if st.button("📑 Export as Word", use_container_width=True):
            doc_bytes = export_to_word(results)
            st.download_button(
                label="💾 Download Word File",
                data=doc_bytes,
                file_name=f"whisperforge_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    with col3:
        if st.button("📋 Copy All Content", use_container_width=True):
            all_content = create_text_export(results)
            st.code(all_content, language="text")
            st.success("✅ Content displayed above - use your browser's copy function!")
        if st.button("📕 Export as PDF", use_container_width=True):
            pdf_bytes = export_to_pdf(results)
            st.download_button(
                label="💾 Download PDF",
                data=pdf_bytes,
                file_name=f"whisperforge_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )

def create_text_export(results):
    """Create a formatted text export of all content"""
    export_lines = []
    export_lines.append("=" * 60)
    export_lines.append("WHISPERFORGE CONTENT EXPORT")
    export_lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    export_lines.append("=" * 60)
    export_lines.append("")
    
    sections = [
        ("AUDIO TRANSCRIPT", results.get('transcript', '')),
        ("EXTRACTED WISDOM", results.get('wisdom', '')),
        ("CONTENT OUTLINE", results.get('outline', '')),
        ("FULL ARTICLE", results.get('article', '')),
        ("SOCIAL MEDIA CONTENT", results.get('social_content', ''))
    ]
    
    for title, content in sections:
        if content:
            export_lines.append(f"## {title}")
            export_lines.append("-" * 40)
            export_lines.append(content)
            export_lines.append("")
            export_lines.append("")
    
    if results.get('notion_url'):
        export_lines.append("## NOTION LINK")
        export_lines.append("-" * 40)
        export_lines.append(results['notion_url'])
        export_lines.append("")

    return "\n".join(export_lines)

def export_to_markdown(results):
    """Export results to Markdown format"""
    lines = ["# WhisperForge Content Export"]
    sections = [
        ("## Transcript", results.get('transcript', '')),
        ("## Wisdom", results.get('wisdom', '')),
        ("## Outline", results.get('outline', '')),
        ("## Article", results.get('article', '')),
        ("## Social Content", results.get('social_content', '')),
    ]
    for title, content in sections:
        if content:
            lines.append(title)
            lines.append("")
            lines.append(content)
            lines.append("")
    if results.get('notion_url'):
        lines.append("## Notion Link")
        lines.append(results['notion_url'])
    return "\n".join(lines)

def export_to_word(results):
    """Export results to a Word document"""
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_heading("WhisperForge Content Export", level=1)
    for title, content in [
        ("Transcript", results.get('transcript', '')),
        ("Wisdom", results.get('wisdom', '')),
        ("Outline", results.get('outline', '')),
        ("Article", results.get('article', '')),
        ("Social Content", results.get('social_content', '')),
    ]:
        if content:
            doc.add_heading(title, level=2)
            doc.add_paragraph(content)
    if results.get('notion_url'):
        doc.add_heading("Notion Link", level=2)
        doc.add_paragraph(results['notion_url'])

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def export_to_pdf(results):
    """Export results to a PDF file"""
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    text = create_text_export(results)
    for line in text.split("\n"):
        pdf.cell(0, 10, txt=line, ln=1)
    return pdf.output(dest='S').encode('latin-1')

# === NAVIGATION & PAGES ===
def create_aurora_navigation():
    """Beautiful Aurora bioluminescent navigation - Clean and professional"""
    st.markdown("""
    <div class="aurora-header">
        <div class="aurora-glow"></div>
        <h1 class="aurora-title">WhisperForge Aurora</h1>
        <p class="aurora-subtitle">Transform Audio into Structured Content with AI</p>
        <div class="aurora-pipeline">
            <span class="pipeline-badge">Transcription</span>
            <span class="pipeline-badge">Wisdom</span>
            <span class="pipeline-badge">Article</span>
            <span class="pipeline-badge">Social</span>
            <span class="pipeline-badge">Notion</span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Clean navigation tabs without emojis
    tabs = st.tabs([
        "Transform", 
        "Content Library", 
        "Settings", 
        "Knowledge Base",
        "Prompts"
    ])
    
    return tabs

def show_transform_page():
    """Clean transformation page focused on file upload and processing"""
    
    # Simple Aurora-styled header using main CSS
    st.markdown("""
    <div class="transform-header">
        <h2 class="transform-title">Transform Audio</h2>
        <p class="transform-subtitle">Upload your audio and watch it transform into structured content</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Beautiful Aurora upload method selector
    st.markdown("""
    <div class="aurora-upload-method-selector">
        <div class="aurora-upload-method-card" id="standard-upload">
            <span class="aurora-upload-method-icon">⚡</span>
            <h3 class="aurora-upload-method-title">Standard Upload</h3>
            <p class="aurora-upload-method-description">Perfect for most audio files up to 25MB with instant processing</p>
            <div class="aurora-upload-method-features">
                <div class="aurora-upload-feature">
                    <span class="aurora-upload-feature-icon">⚡</span>
                    <span>Instant processing</span>
                </div>
                <div class="aurora-upload-feature">
                    <span class="aurora-upload-feature-icon">🎵</span>
                    <span>Audio preview</span>
                </div>
                <div class="aurora-upload-feature">
                    <span class="aurora-upload-feature-icon">📊</span>
                    <span>Up to 25MB</span>
                </div>
            </div>
        </div>
        <div class="aurora-upload-method-card" id="large-upload">
            <span class="aurora-upload-method-icon">🚀</span>
            <h3 class="aurora-upload-method-title">Large File Upload</h3>
            <p class="aurora-upload-method-description">Advanced processing for large files up to 2GB with intelligent chunking</p>
            <div class="aurora-upload-method-features">
                <div class="aurora-upload-feature">
                    <span class="aurora-upload-feature-icon">🔧</span>
                    <span>FFmpeg chunking</span>
                </div>
                <div class="aurora-upload-feature">
                    <span class="aurora-upload-feature-icon">⚡</span>
                    <span>Parallel processing</span>
                </div>
                <div class="aurora-upload-feature">
                    <span class="aurora-upload-feature-icon">📈</span>
                    <span>Up to 2GB</span>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    document.addEventListener('DOMContentLoaded', function() {
        // Add click handlers to upload method cards
        const standardCard = document.getElementById('standard-upload');
        const largeCard = document.getElementById('large-upload');
        
        if (standardCard && largeCard) {
            // Initialize with first option selected
            standardCard.classList.add('selected');
            
            standardCard.addEventListener('click', function() {
                // Remove selected class from both cards
                standardCard.classList.remove('selected');
                largeCard.classList.remove('selected');
                
                // Add selected class to clicked card
                standardCard.classList.add('selected');
                
                // Find and trigger the radio button
                setTimeout(() => {
                    const radioButtons = document.querySelectorAll('input[type="radio"]');
                    radioButtons.forEach(radio => {
                        if (radio.nextElementSibling && radio.nextElementSibling.textContent.includes('Standard Upload')) {
                            radio.checked = true;
                            radio.dispatchEvent(new Event('change', { bubbles: true }));
                        }
                    });
                }, 100);
            });
            
            largeCard.addEventListener('click', function() {
                // Remove selected class from both cards
                standardCard.classList.remove('selected');
                largeCard.classList.remove('selected');
                
                // Add selected class to clicked card
                largeCard.classList.add('selected');
                
                // Find and trigger the radio button
                setTimeout(() => {
                    const radioButtons = document.querySelectorAll('input[type="radio"]');
                    radioButtons.forEach(radio => {
                        if (radio.nextElementSibling && radio.nextElementSibling.textContent.includes('Large File Upload')) {
                            radio.checked = true;
                            radio.dispatchEvent(new Event('change', { bubbles: true }));
                        }
                    });
                }, 100);
            });
        }
    });
    </script>
    """, unsafe_allow_html=True)
    
    # Enhanced file upload selection with session state
    if 'upload_method' not in st.session_state:
        st.session_state.upload_method = "Standard Upload"
    
    upload_method = st.radio(
        "Choose upload method:",
        ["Standard Upload", "Large File Upload"],
        index=0 if st.session_state.upload_method == "Standard Upload" else 1,
        help="Standard upload for smaller files, Enhanced upload for large files with FFmpeg processing",
        label_visibility="collapsed"
    )
    
    st.session_state.upload_method = upload_method
    
    if upload_method == "Standard Upload":
        # Beautiful standard file upload zone
        st.markdown("""
        <div class="aurora-file-upload-zone">
            <div class="aurora-upload-icon-container">
                <div class="aurora-upload-icon">🎵</div>
                <div class="aurora-upload-pulse"></div>
            </div>
            <h3 class="aurora-upload-title">Drop your audio file here</h3>
            <p class="aurora-upload-subtitle">Or click to browse and select a file</p>
            <div class="aurora-upload-formats">
                <span class="aurora-format-badge">MP3</span>
                <span class="aurora-format-badge">WAV</span>
                <span class="aurora-format-badge">M4A</span>
                <span class="aurora-format-badge">FLAC</span>
                <span class="aurora-format-badge">OGG</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Standard file upload
        uploaded_files = st.file_uploader(
            "Upload your audio file",
            type=['mp3', 'wav', 'm4a', 'flac', 'ogg'],
            help="Upload audio file for processing (max 25MB)",
            label_visibility="collapsed",
            accept_multiple_files=True
        )

        if uploaded_files:
            for uploaded_file in uploaded_files:
                # Beautiful file preview card
                file_size = len(uploaded_file.getvalue()) / (1024 * 1024)
                file_extension = uploaded_file.name.split('.')[-1].upper()
            
            st.markdown(f"""
            <div class="aurora-file-preview">
                <div class="aurora-file-preview-header">
                    <div class="aurora-file-info">
                        <div class="aurora-file-icon">🎵</div>
                        <div class="aurora-file-details">
                            <h4>{uploaded_file.name}</h4>
                            <p>{file_size:.1f} MB • {file_extension} Format</p>
                        </div>
                    </div>
                    <div class="aurora-file-actions">
                        <div class="aurora-file-action-btn">Ready to process</div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
                # Enhanced audio player
                if file_size < 50:  # Only show player for files under 50MB
                    st.markdown('<div class="aurora-audio-player">', unsafe_allow_html=True)
                    st.audio(uploaded_file.getvalue())
                    st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.info("Audio preview disabled for large files to conserve memory")

                # Beautiful process button
                if st.button(f"Transform {uploaded_file.name}", key=f"process_{uploaded_file.name}", type="primary", use_container_width=True):
                    if not os.getenv("OPENAI_API_KEY"):
                        st.error("Please enter your OpenAI API key in the sidebar")
                        return

                    with st.container():
                        if st.session_state.get('live_stream', False):
                            results = process_audio_pipeline_live(uploaded_file)
                        else:
                            results = process_audio_pipeline(uploaded_file)
                        if results:
                            st.session_state.current_results = results
    
    else:
        # Enhanced large file upload
        st.markdown("### Enhanced Large File Processing")
        
        # Initialize enhanced processor
        processor = EnhancedLargeFileProcessor()
        
        # Create enhanced upload interface
        uploaded_file = processor.create_enhanced_upload_interface()
        
        if uploaded_file:
            # Validate file
            validation = processor.validate_file(uploaded_file)
            
            if not validation["valid"]:
                st.error(f"{validation['error']}")
                return
            
            file_size_mb = validation["size_mb"]
            requires_chunking = validation["requires_chunking"]
            
            # Beautiful processing metrics
            st.markdown(f"""
            <div class="aurora-processing-metrics">
                <div class="aurora-metric-card">
                    <div class="aurora-metric-value">{file_size_mb:.1f}</div>
                    <div class="aurora-metric-label">File Size (MB)</div>
                </div>
                <div class="aurora-metric-card">
                    <div class="aurora-metric-value">{"FFmpeg" if requires_chunking else "Standard"}</div>
                    <div class="aurora-metric-label">Processing Method</div>
                </div>
                <div class="aurora-metric-card">
                    <div class="aurora-metric-value">{validation["format"].upper()}</div>
                    <div class="aurora-metric-label">Format</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # Audio preview disabled for large files to conserve memory
            if file_size_mb < 50:
                st.markdown('<div class="aurora-audio-player">', unsafe_allow_html=True)
                st.audio(uploaded_file.getvalue())
                st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.info("Audio preview disabled for large files to conserve memory")
            
            # Enhanced process button
            if st.button("Transform Large File to Content", type="primary", use_container_width=True):
                if not os.getenv("OPENAI_API_KEY"):
                    st.error("Please enter your OpenAI API key in the sidebar")
                    return
                
                with st.container():
                    # Process with enhanced large file processor
                    processing_result = processor.process_large_file(uploaded_file)
                    
                    if processing_result["success"]:
                        transcript = processing_result["transcript"]
                        
                        # Show processing summary
                        st.success(f"Large file processing complete!")
                        
                        # Beautiful success metrics
                        st.markdown(f"""
                        <div class="aurora-processing-metrics">
                            <div class="aurora-metric-card">
                                <div class="aurora-metric-value">{len(transcript):,}</div>
                                <div class="aurora-metric-label">Characters</div>
                            </div>
                            <div class="aurora-metric-card">
                                <div class="aurora-metric-value">{processing_result["method"]}</div>
                                <div class="aurora-metric-label">Method Used</div>
                            </div>
                            <div class="aurora-metric-card">
                                <div class="aurora-metric-value">{processing_result.get("chunks_processed", 1)}</div>
                                <div class="aurora-metric-label">Chunks Processed</div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Continue with pipeline using pre-transcribed content
                        st.markdown("---")
                        results = process_audio_pipeline_with_transcript(transcript)
                        
                        if results:
                            # Store results in session state
                            st.session_state.current_results = results
                    else:
                        st.error(f"Large file processing failed: {processing_result['error']}")
                        
                        # Fallback to standard processing for smaller files
                        if file_size_mb < 100:
                            st.info("Attempting fallback to standard processing...")
                            try:
                                if st.session_state.get('live_stream', False):
                                    results = process_audio_pipeline_live(uploaded_file)
                                else:
                                    results = process_audio_pipeline(uploaded_file)
                                if results:
                                    st.session_state.current_results = results
                            except Exception as e:
                                st.error(f"Fallback processing also failed: {str(e)}")
    
    # Show results if available
    if 'current_results' in st.session_state:
        show_results(st.session_state.current_results)

def show_content_library():
    """Content library/history page"""
    st.markdown("### 📚 Content Library")
    
    # Get content from database
    try:
        db = get_supabase_client()
        if db:
            # Fetch recent content
            response = db.client.table('content').select('*').order('created_at', desc=True).limit(20).execute()
            
            if response.data:
                st.success(f"📊 Found {len(response.data)} content items")
                
                # Search and filter
                col1, col2 = st.columns([3, 1])
                with col1:
                    search_term = st.text_input("🔍 Search content", placeholder="Search by title or content...")
                with col2:
                    content_type = st.selectbox("Filter by type", ["All", "Article", "Social", "Outline"])
                
                # Display content cards in a two-column grid
                cols = st.columns(2)
                for idx, item in enumerate(response.data):
                    if search_term and search_term.lower() not in item.get('title', '').lower():
                        continue

                    with cols[idx % 2].expander(f"📄 {item.get('title', 'Untitled')} - {item.get('created_at', '')[:10]}"):
                        col1, col2 = st.columns([3, 1])
                        
                        with col1:
                            st.markdown(f"**Created:** {item.get('created_at', 'Unknown')}")
                            if item.get('transcript'):
                                st.markdown("**Transcript Preview:**")
                                st.text(item['transcript'][:200] + "..." if len(item['transcript']) > 200 else item['transcript'])
                        
                        with col2:
                            if st.button(f"🔄 Reprocess", key=f"reprocess_{item.get('id')}"):
                                st.info("Reprocessing feature coming soon!")
                            
                            if st.button(f"📤 Export", key=f"export_{item.get('id')}"):
                                st.info("Export feature coming soon!")
                        
                        # Show generated content
                        if item.get('wisdom'):
                            st.markdown("**💡 Wisdom:**")
                            st.text_area("", item['wisdom'], height=100, disabled=True, key=f"wisdom_{item.get('id')}")
                        
                        if item.get('article'):
                            st.markdown("**📰 Article:**") 
                            st.text_area("", item['article'], height=150, disabled=True, key=f"article_{item.get('id')}")
                        
                        if item.get('social_content'):
                            st.markdown("**📱 Social Content:**")
                            st.text_area("", item['social_content'], height=100, disabled=True, key=f"social_{item.get('id')}")
            else:
                st.info("📭 No content found. Process some audio files to see them here!")
        else:
            st.error("❌ Database connection failed")
    except Exception as e:
        st.error(f"❌ Error loading content library: {e}")

def show_settings_page():
    """Settings and configuration page"""
    st.markdown("### ⚙️ Settings & Configuration")
    
    # API Keys section
    st.markdown("#### 🔑 API Keys")
    with st.expander("🔧 API Configuration", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            # OpenAI settings
            st.markdown("**OpenAI Configuration**")
            openai_key = st.text_input("OpenAI API Key", type="password", 
                                      value=os.getenv("OPENAI_API_KEY", ""),
                                      help="Your OpenAI API key")
            if openai_key:
                os.environ["OPENAI_API_KEY"] = openai_key
                st.success("✅ OpenAI key configured")
            
            # Model selection
            model_choice = st.selectbox("OpenAI Model", 
                                       ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo"],
                                       help="Choose the OpenAI model for content generation")
            st.session_state.openai_model = model_choice
        
        with col2:
            # Notion settings
            st.markdown("**Notion Configuration**")
            notion_key = st.text_input("Notion API Key", type="password",
                                      value=os.getenv("NOTION_API_KEY", ""),
                                      help="Your Notion integration token")
            if notion_key:
                os.environ["NOTION_API_KEY"] = notion_key
            
            notion_db = st.text_input("Notion Database ID",
                                     value=os.getenv("NOTION_DATABASE_ID", ""),
                                     help="Your Notion database ID")
            if notion_db:
                os.environ["NOTION_DATABASE_ID"] = notion_db
            
            if notion_key and notion_db:
                st.success("✅ Notion configured")
    
    # Pipeline settings
    st.markdown("#### 🔄 Pipeline Configuration")
    with st.expander("⚙️ Processing Pipeline", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Core Features**")
            auto_notion = st.checkbox("Auto-publish to Notion",
                                     value=st.session_state.get('auto_notion', True))
            st.session_state.auto_notion = auto_notion

            live_stream = st.checkbox(
                "Live Streaming",
                value=st.session_state.get('live_stream', False),
                help="Show step-by-step streaming results"
            )
            st.session_state.live_stream = live_stream
            
            large_file_mode = st.checkbox("Enhanced Large File Processing", 
                                        value=st.session_state.get('large_file_mode', True),
                                        help="Use FFmpeg for files larger than 25MB")
            st.session_state.large_file_mode = large_file_mode
        
        with col2:
            st.markdown("**Quality Settings**")
            content_length = st.selectbox("Article Length", 
                                        ["Short (500-800 words)", "Medium (800-1200 words)", "Long (1200+ words)"])
            
            tone_style = st.selectbox("Content Tone",
                                    ["Professional", "Conversational", "Academic", "Creative"])

            st.session_state.content_length = content_length
            st.session_state.tone_style = tone_style

            templates = [f.replace('.md','') for f in os.listdir('templates')] if os.path.exists('templates') else []
            if templates:
                template_choice = st.selectbox("Article Template", templates)
                st.session_state.article_template = template_choice
            else:
                st.session_state.article_template = None
    
    # System status
    st.markdown("#### 🔍 System Status")
    with st.expander("📊 Connection Status", expanded=False):
        if st.button("🧪 Test All Connections"):
            with st.spinner("Testing all connections..."):
                # Test OpenAI
                try:
                    if os.getenv("OPENAI_API_KEY"):
                        st.success("✅ OpenAI API key configured")
                    else:
                        st.error("❌ OpenAI API key missing")
                except Exception as e:
                    st.error(f"❌ OpenAI error: {e}")
                
                # Test Supabase
                try:
                    db = get_supabase_client()
                    if db and db.test_connection():
                        st.success("✅ Supabase connected")
                    else:
                        st.error("❌ Supabase connection failed")
                except Exception as e:
                    st.error(f"❌ Supabase error: {e}")
                
                # Test Notion
                try:
                    if os.getenv("NOTION_API_KEY") and os.getenv("NOTION_DATABASE_ID"):
                        from notion_client import Client
                        client = Client(auth=os.getenv("NOTION_API_KEY"))
                        database = client.databases.retrieve(database_id=os.getenv("NOTION_DATABASE_ID"))
                        st.success("✅ Notion connected")
                    else:
                        st.warning("⚠️ Notion not configured")
                except Exception as e:
                    st.error(f"❌ Notion error: {e}")

def show_knowledge_base():
    """Knowledge base management page"""
    st.markdown("### 🧠 Knowledge Base")
    
    # Check if knowledge base files exist
    kb_path = "prompts/default/knowledge_base"
    
    st.markdown("""
    The knowledge base provides context and expertise to enhance content generation.
    Add domain-specific information, style guides, and reference materials here.
    """)
    
    # Knowledge base sections
    tabs = st.tabs(["📖 View Knowledge", "➕ Add Knowledge", "🔧 Manage Files"])
    
    with tabs[0]:
        st.markdown("#### 📖 Current Knowledge Base")
        try:
            if os.path.exists(kb_path):
                files = [f for f in os.listdir(kb_path) if f.endswith('.md')]
                if files:
                    selected_file = st.selectbox("Select knowledge file:", files)
                    if selected_file:
                        file_path = os.path.join(kb_path, selected_file)
                        with open(file_path, 'r') as f:
                            content = f.read()
                        
                        st.markdown(f"**File:** `{selected_file}`")
                        create_enhanced_aurora_content_card("Knowledge Content", content, "text", "📖")
                else:
                    st.info("📭 No knowledge files found")
            else:
                st.info("📁 Knowledge base directory not found")
        except Exception as e:
            st.error(f"❌ Error reading knowledge base: {e}")
    
    with tabs[1]:
        st.markdown("#### ➕ Add New Knowledge")
        
        col1, col2 = st.columns([2, 1])
        with col1:
            kb_title = st.text_input("Knowledge Title", placeholder="e.g., 'Marketing Guidelines'")
        with col2:
            kb_category = st.selectbox("Category", ["General", "Style Guide", "Domain Expertise", "Templates"])
        
        kb_content = st.text_area("Knowledge Content", 
                                 placeholder="Enter your knowledge content here...",
                                 height=300)
        
        if st.button("💾 Save Knowledge", type="primary"):
            if kb_title and kb_content:
                try:
                    os.makedirs(kb_path, exist_ok=True)
                    filename = f"{kb_title.lower().replace(' ', '_')}.md"
                    file_path = os.path.join(kb_path, filename)
                    
                    with open(file_path, 'w') as f:
                        f.write(f"# {kb_title}\n\n")
                        f.write(f"**Category:** {kb_category}\n\n")
                        f.write(kb_content)
                    
                    st.success(f"✅ Knowledge saved as `{filename}`")
                except Exception as e:
                    st.error(f"❌ Error saving knowledge: {e}")
            else:
                st.error("❌ Please provide both title and content")
    
    with tabs[2]:
        st.markdown("#### 🔧 Manage Knowledge Files")
        
        try:
            if os.path.exists(kb_path):
                files = [f for f in os.listdir(kb_path) if f.endswith('.md')]
                if files:
                    for file in files:
                        col1, col2, col3 = st.columns([3, 1, 1])
                        with col1:
                            st.markdown(f"📄 `{file}`")
                        with col2:
                            if st.button("📝 Edit", key=f"edit_{file}"):
                                st.info("Edit functionality coming soon!")
                        with col3:
                            if st.button("🗑️ Delete", key=f"delete_{file}"):
                                try:
                                    os.remove(os.path.join(kb_path, file))
                                    st.success(f"✅ Deleted `{file}`")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Error deleting file: {e}")
                else:
                    st.info("📭 No knowledge files found")
            else:
                st.info("📁 Knowledge base directory not found")
        except Exception as e:
            st.error(f"❌ Error managing files: {e}")

def show_prompts_page():
    """Prompts management page"""
    st.markdown("### 📝 Prompts Management")
    
    st.markdown("""
    Customize the AI prompts used in each step of the content generation pipeline.
    Fine-tune the output style, format, and focus for your specific needs.
    """)
    
    # Prompt categories
    prompt_types = {
        "wisdom": "💡 Wisdom Extraction",
        "outline": "📋 Content Outline", 
        "article": "📰 Article Generation",
        "social": "📱 Social Media Posts"
    }
    
    tabs = st.tabs(list(prompt_types.values()) + ["🔧 Advanced"])
    
    for i, (prompt_key, prompt_name) in enumerate(prompt_types.items()):
        with tabs[i]:
            st.markdown(f"#### {prompt_name}")
            
            # Load current prompt - Map UI keys to actual pipeline files
            file_mapping = {
                "wisdom": "wisdom_extraction.md",
                "outline": "outline_creation.md", 
                "article": "article_generation.md",
                "social": "social_media.md"
            }
            prompt_file = f"prompts/default/{file_mapping[prompt_key]}"
            current_prompt = ""
            
            try:
                if os.path.exists(prompt_file):
                    with open(prompt_file, 'r') as f:
                        current_prompt = f.read()
                else:
                    current_prompt = f"# {prompt_name} Prompt\n\nDefault prompt for {prompt_key} generation."
            except Exception as e:
                st.error(f"❌ Error loading prompt: {e}")
            
            # Edit prompt
            new_prompt = st.text_area(
                f"Edit {prompt_name} Prompt",
                value=current_prompt,
                height=400,
                help=f"Customize the prompt used for {prompt_key} generation"
            )
            
            col1, col2, col3 = st.columns([1, 1, 2])
            with col1:
                if st.button(f"💾 Save", key=f"save_{prompt_key}"):
                    try:
                        os.makedirs("prompts/default", exist_ok=True)
                        with open(prompt_file, 'w') as f:
                            f.write(new_prompt)
                        st.success(f"✅ {prompt_name} prompt saved!")
                    except Exception as e:
                        st.error(f"❌ Error saving prompt: {e}")
            
            with col2:
                if st.button(f"🔄 Reset", key=f"reset_{prompt_key}"):
                    st.info("Reset to default functionality coming soon!")
            
            with col3:
                st.markdown(f"**File:** `{prompt_file}`")
    
    # Advanced settings
    with tabs[-1]:
        st.markdown("#### 🔧 Advanced Prompt Settings")
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Global Settings**")
            temperature = st.slider("Temperature (Creativity)", 0.0, 1.0, 0.7, 0.1)
            max_tokens = st.number_input("Max Tokens", 100, 4000, 2000)
            
        with col2:
            st.markdown("**Prompt Templates**")
            if st.button("📥 Import Prompt Set"):
                st.info("Import functionality coming soon!")
            if st.button("📤 Export Prompt Set"):
                st.info("Export functionality coming soon!")
        
        st.session_state.temperature = temperature
        st.session_state.max_tokens = max_tokens

def create_aurora_tabs(tab_data, default_tab=0):
    """Create beautiful Aurora-styled tabs with a simplified, reliable approach"""
    import uuid
    
    # Generate unique ID for this tab group
    tab_group_id = f"tabs_{uuid.uuid4().hex[:8]}"
    
    # Initialize session state for this tab group
    if f"{tab_group_id}_active" not in st.session_state:
        st.session_state[f"{tab_group_id}_active"] = default_tab
    
    # Calculate content stats for display
    tab_stats = []
    for tab in tab_data:
        content = tab.get('content', '')
        word_count = len(str(content).split()) if content else 0
        tab_stats.append(word_count)
    
    # Create the tabs container with Aurora styling
    st.markdown("""
    <div class="aurora-tabs-container">
        <div class="aurora-tabs-actions">
            <h3 class="aurora-tabs-title">
                <span>📄</span>
                Generated Content
            </h3>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Create tab selector using Streamlit's selectbox with Aurora styling
    tab_options = []
    for i, tab in enumerate(tab_data):
        word_count = tab_stats[i]
        tab_options.append(f"{tab['icon']} {tab['title']} ({word_count} words)")
    
    # Custom styled selectbox
    st.markdown("""
    <style>
    .aurora-tab-selector .stSelectbox > div > div {
        background: var(--aurora-bg-glass) !important;
        border: 2px solid var(--aurora-border) !important;
        border-radius: var(--aurora-radius) !important;
        color: var(--aurora-text) !important;
    }
    .aurora-tab-selector .stSelectbox > div > div:hover {
        border-color: var(--aurora-border-hover) !important;
        box-shadow: var(--aurora-glow-subtle) !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Tab selector
    with st.container():
        st.markdown('<div class="aurora-tab-selector">', unsafe_allow_html=True)
        selected_tab_label = st.selectbox(
            "Select Content Type:",
            tab_options,
            index=st.session_state[f"{tab_group_id}_active"],
            key=f"tab_select_{tab_group_id}",
            label_visibility="collapsed"
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Find selected tab index
    selected_index = 0
    for i, option in enumerate(tab_options):
        if option == selected_tab_label:
            selected_index = i
            break
    
    # Update session state
    st.session_state[f"{tab_group_id}_active"] = selected_index
    
    # Display selected content with Aurora styling
    if 0 <= selected_index < len(tab_data):
        active_tab = tab_data[selected_index]
        
        # Add quick actions for the active tab
        col1, col2, col3, col4 = st.columns([1, 1, 1, 2])
        
        with col1:
            if st.button("📋 Copy", key=f"copy_{tab_group_id}_{selected_index}", use_container_width=True):
                st.code(active_tab.get('content', ''), language='text')
                st.success("✅ Content displayed above - copy with Ctrl+A, Ctrl+C")
        
        with col2:
            content = active_tab.get('content', '')
            if content:
                st.download_button(
                    label="💾 Download",
                    data=content,
                    file_name=f"{active_tab['title'].lower().replace(' ', '_')}.txt",
                    mime="text/plain",
                    key=f"download_{tab_group_id}_{selected_index}",
                    use_container_width=True
                )
        
        with col3:
            if st.button("📊 Stats", key=f"stats_{tab_group_id}_{selected_index}", use_container_width=True):
                word_count = len(str(content).split())
                char_count = len(str(content))
                st.info(f"📊 **{active_tab['title']}**: {word_count} words, {char_count} characters")
        
        # Display the content using our enhanced content card
        st.markdown('<div class="aurora-tab-content">', unsafe_allow_html=True)
        create_enhanced_aurora_content_card(
            title=active_tab['title'],
            content=active_tab.get('content', ''),
            content_type=active_tab.get('type', 'text'),
            icon=active_tab['icon']
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    return selected_index

# === MAIN APP ===
def show_main_app():
    """Main application interface with navigation"""
    # Create navigation
    tabs = create_aurora_navigation()
    
    # Show different pages based on selected tab
    with tabs[0]:  # Transform
        show_transform_page()
    
    with tabs[1]:  # Content Library
        show_content_library()
    
    with tabs[2]:  # Settings
        show_settings_page()
    
    with tabs[3]:  # Knowledge Base
        show_knowledge_base()
    
    with tabs[4]:  # Prompts
        show_prompts_page()

# === ENTRY POINT ===
def main():
    """Application entry point"""
    init_session()
    
    if st.session_state.authenticated:
        show_main_app()
    else:
        show_login()

if __name__ == "__main__":
    main() 